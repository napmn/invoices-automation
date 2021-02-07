import argparse
import json
import os
import re
from datetime import datetime
from glob import glob
from functools import lru_cache

from docx import Document
from docx.text.paragraph import Paragraph


@lru_cache(maxsize=1)
def get_variables():
    args = get_cmd_args()
    return {
        'invoice_id': get_next_invoice_id(),
        'invoice_issue_date': args['date'],
        'mandays': args['mandays'],
        'total': lambda variables: variables['mandays'] * variables['md_rate']
    }


@lru_cache(maxsize=1)
def get_cmd_args():
    parser = argparse.ArgumentParser()
    subparsers = parser.add_subparsers(dest='command', required=True)
    subparsers.add_parser(
        'print',
        description='Prints paragraphs of the invoice document',
        help='Prints paragraphs of invoice'
    )
    parser_for_create = subparsers.add_parser(
        'create',
        description='Creates new invoice based on older invoice',
        help='Creates new invoice'
    )
    parser_for_create.add_argument(
        '--mandays', type=int, required=True, help='Number of worked "man days"'
    )
    parser_for_create.add_argument(
        '--date', type=str, required=False,
        default='{0.day}. {0.month}. {0.year}'.format(datetime.today()),
        help='Invoice issue date'
    )
    return vars(parser.parse_args())


@lru_cache(maxsize=1)
def load_config() -> dict:
    with open('invoice_config.json', 'r') as config_f:
        return json.load(config_f)


def print_paragraphs(doc: Document):
    for i, paragraph in enumerate(doc.paragraphs):
        print(f'{i}. {paragraph.text}')


def replace_variable_value(p: Paragraph, var_name: str):
    config = load_config()
    variables = get_variables()

    pattern = config['regexes'][var_name]
    for run in p.runs:
        # in case there are multiple Runs
        previous_text = run.text
        if re.search(pattern, previous_text):
            # first check if value is not in config
            new_val = config['constants'].get(var_name, None)
            new_val = variables[var_name] if new_val is None else new_val
            run.clear()  # clear text in Run but preserve style
            if callable(new_val):
                new_val = new_val(variables | config['constants'])
            run.add_text(re.sub(pattern, str(new_val), previous_text))
            break


def get_next_invoice_id():
    current_year = datetime.today().year
    base = config['general']['invoice_base_name'] 
    glob_pattern = os.path.join(
        config['general']['output_folder_path'],
        f'{base}{current_year}*.docx'
    )
    invoices = glob(glob_pattern)
    latest_invoice_number = max([int(i[-9:-5]) for i in invoices]) if invoices else 0
    return f'{current_year}{latest_invoice_number + 1:04}'
    

def create_invoice(doc: Document):
    config = load_config()
    variables = get_variables()

    for p_num, data in config['paragraphs'].items():
        p = doc.paragraphs[int(p_num)]  # get paragraph that needs to be changed
        for var_name in data['variables']:
            replace_variable_value(p, var_name)
 
    output_path = os.path.join(
        config['general']['output_folder_path'],
        f'{config["general"]["invoice_base_name"]}{variables["invoice_id"]}.docx'
    )
    print(f'Saving invoice to {output_path}')
    doc.save(output_path)


if __name__ == '__main__':
    args = get_cmd_args()
    config = load_config()

    commands = {'print': print_paragraphs, 'create': create_invoice}
    doc = Document(config['general']['invoice_template_path'])
    commands[args['command']](doc)

